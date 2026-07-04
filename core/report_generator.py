from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import plotly.io as pio

from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors as rl_colors

ACCENT_COLOR = RGBColor(0xD9, 0x7A, 0x3F)  # matches our theme's copper accent

def _add_heading(doc, text, size=16):
    heading = doc.add_paragraph()
    run = heading.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = ACCENT_COLOR
    return heading

def _fig_to_image_bytes(fig):
    """Converts a Plotly figure to PNG bytes so it can be embedded in a document."""
    img_bytes = pio.to_image(fig, format="png", width=700, height=400, scale=2)
    return io.BytesIO(img_bytes)

def generate_report(df, summary, insights, charts, chat_history, filename="dataset"):
    """
    Builds a Word document summarizing the dataset, insights, charts,
    and the Q&A conversation. Returns a BytesIO object ready for download.
    """
    doc = Document()

    title = doc.add_paragraph()
    title_run = title.add_run("AutoInsight Analysis Report")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = ACCENT_COLOR
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run(f"Dataset: {filename}")
    subtitle_run.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    _add_heading(doc, "Dataset Overview")
    doc.add_paragraph(f"Rows: {summary['num_rows']:,}")
    doc.add_paragraph(f"Columns: {summary['num_columns']}")
    doc.add_paragraph(f"Missing values: {sum(summary['missing_values'].values())}")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 4"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Column"
    hdr_cells[1].text = "Type"
    hdr_cells[2].text = "Missing"
    for col in summary["columns"]:
        row_cells = table.add_row().cells
        row_cells[0].text = col
        row_cells[1].text = summary["dtypes"][col]
        row_cells[2].text = str(summary["missing_values"][col])

    doc.add_paragraph()

    _add_heading(doc, "Automatic Insights")
    for insight in insights:
        doc.add_paragraph(insight, style="List Bullet")

    doc.add_paragraph()

    _add_heading(doc, "Visualizations")
    for title_text, fig in charts:
        doc.add_paragraph(title_text, style="Intense Quote")
        try:
            img_stream = _fig_to_image_bytes(fig)
            doc.add_picture(img_stream, width=Inches(6))
        except Exception as e:
            doc.add_paragraph(f"[Chart could not be rendered: {e}]")
        doc.add_paragraph()

    if chat_history:
        _add_heading(doc, "Questions & Answers")
        for msg in chat_history:
            p = doc.add_paragraph()
            role_label = "Q: " if msg["role"] == "user" else "A: "
            run = p.add_run(role_label)
            run.bold = True
            p.add_run(msg["content"])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def generate_pdf_report(df, summary, insights, charts, chat_history, filename="dataset"):
    """Builds a PDF version of the same report using reportlab."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.6*inch, bottomMargin=0.6*inch)
    styles = getSampleStyleSheet()

    accent = rl_colors.HexColor("#D97A3F")
    title_style = ParagraphStyle("TitleStyle", parent=styles["Title"], textColor=accent, fontSize=22)
    heading_style = ParagraphStyle("HeadingStyle", parent=styles["Heading2"], textColor=accent, spaceBefore=14)

    story = []
    story.append(Paragraph("AutoInsight Analysis Report", title_style))
    story.append(Paragraph(f"Dataset: {filename}", styles["Italic"]))
    story.append(Spacer(1, 16))

    story.append(Paragraph("Dataset Overview", heading_style))
    story.append(Paragraph(
        f"Rows: {summary['num_rows']:,} &nbsp;&nbsp; Columns: {summary['num_columns']} &nbsp;&nbsp; "
        f"Missing values: {sum(summary['missing_values'].values())}",
        styles["Normal"]
    ))
    story.append(Spacer(1, 10))

    table_data = [["Column", "Type", "Missing"]]
    for col in summary["columns"]:
        table_data.append([col, summary["dtypes"][col], str(summary["missing_values"][col])])
    t = Table(table_data, colWidths=[200, 150, 100])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), accent),
        ("TEXTCOLOR", (0, 0), (-1, 0), rl_colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, rl_colors.grey),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
    ]))
    story.append(t)
    story.append(Spacer(1, 16))

    story.append(Paragraph("Automatic Insights", heading_style))
    for insight in insights:
        story.append(Paragraph(f"• {insight}", styles["Normal"]))
    story.append(Spacer(1, 16))

    story.append(Paragraph("Visualizations", heading_style))
    for title_text, fig in charts:
        story.append(Paragraph(title_text, styles["Heading4"]))
        try:
            img_stream = _fig_to_image_bytes(fig)
            story.append(Image(img_stream, width=6.2*inch, height=3.5*inch))
        except Exception as e:
            story.append(Paragraph(f"[Chart could not be rendered: {e}]", styles["Normal"]))
        story.append(Spacer(1, 10))

    if chat_history:
        story.append(Paragraph("Questions & Answers", heading_style))
        for msg in chat_history:
            label = "Q: " if msg["role"] == "user" else "A: "
            story.append(Paragraph(f"<b>{label}</b>{msg['content']}", styles["Normal"]))
            story.append(Spacer(1, 4))

    doc.build(story)
    buffer.seek(0)
    return buffer